"""
Document Processing Module for RAG Chatbot
Handles document loading, chunking, and preprocessing
"""
import os
import logging
from typing import List, Optional
from pathlib import Path

import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

from .config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document types for RAG pipeline"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Ensure directories exist
        os.makedirs(settings.upload_dir, exist_ok=True)
        os.makedirs(settings.processed_dir, exist_ok=True)
    
    def load_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    def load_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def load_txt(self, file_path: str) -> str:
        """Load text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise
    
    def load_html(self, file_path: str) -> str:
        """Extract text from HTML file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text().strip()
        except Exception as e:
            logger.error(f"Error loading HTML {file_path}: {e}")
            raise
    
    def load_document(self, file_path: str) -> str:
        """Load document based on file extension"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        loaders = {
            '.pdf': self.load_pdf,
            '.docx': self.load_docx,
            '.doc': self.load_docx,
            '.txt': self.load_txt,
            '.html': self.load_html,
            '.htm': self.load_html
        }
        
        if extension not in loaders:
            raise ValueError(f"Unsupported file type: {extension}")
        
        logger.info(f"Loading document: {file_path}")
        return loaders[extension](str(file_path))
    
    def chunk_text(self, text: str, metadata: Optional[dict] = None) -> List[LangChainDocument]:
        """Split text into chunks for embedding"""
        try:
            chunks = self.text_splitter.split_text(text)
            documents = []
            
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata or {}
                doc_metadata.update({"chunk_id": i, "chunk_size": len(chunk)})
                
                documents.append(LangChainDocument(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            logger.info(f"Created {len(documents)} chunks")
            return documents
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise
    
    def process_documents(self, file_paths: List[str]) -> List[LangChainDocument]:
        """Process multiple documents and return chunks"""
        all_documents = []
        
        for file_path in file_paths:
            try:
                # Load document
                text = self.load_document(file_path)
                
                # Create metadata
                metadata = {
                    "source": file_path,
                    "filename": Path(file_path).name,
                    "file_type": Path(file_path).suffix
                }
                
                # Chunk the document
                chunks = self.chunk_text(text, metadata)
                all_documents.extend(chunks)
                
                logger.info(f"Processed {file_path}: {len(chunks)} chunks")
                
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue
        
        logger.info(f"Total processed documents: {len(all_documents)}")
        return all_documents


# Example usage
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Example: Process a single document
    # documents = processor.process_documents(["./documents/sample.pdf"])
    # print(f"Processed {len(documents)} document chunks")
